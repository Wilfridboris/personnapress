"""Unit tests for services/ingestion.py."""
import io
import sys
from types import ModuleType
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

# ── Stub heavy optional dependencies not installed in test env ────────────────
# docx is tested separately below


# ── extract_clean_text ────────────────────────────────────────────────────────

def test_extract_clean_text_removes_nav_and_footer():
    from app.services.ingestion import extract_clean_text

    html = """
    <html><body>
      <nav>menu items</nav>
      <main><p>Good content here</p></main>
      <footer>Footer text</footer>
    </body></html>
    """
    result = extract_clean_text(html)
    assert "Good content here" in result
    assert "menu items" not in result
    assert "Footer text" not in result


def test_extract_clean_text_strips_class_patterns():
    from app.services.ingestion import extract_clean_text

    html = """
    <html><body>
      <div class="cookie-banner">Accept cookies</div>
      <div class="sidebar ad-block">Buy stuff</div>
      <article><p>Real content</p></article>
    </body></html>
    """
    result = extract_clean_text(html)
    assert "Real content" in result
    assert "Accept cookies" not in result
    assert "Buy stuff" not in result


def test_extract_clean_text_handles_empty_body():
    from app.services.ingestion import extract_clean_text

    result = extract_clean_text("<html><body></body></html>")
    assert isinstance(result, str)


def test_extract_clean_text_collapses_blank_lines():
    from app.services.ingestion import extract_clean_text

    # Three or more consecutive newlines should be collapsed to at most two
    html = "<main><p>Paragraph one</p><p></p><p></p><p></p><p>Paragraph two</p></main>"
    result = extract_clean_text(html)
    # Should not have three+ consecutive newlines
    assert "\n\n\n" not in result
    assert "Paragraph one" in result
    assert "Paragraph two" in result


# ── extract_file_text ─────────────────────────────────────────────────────────

def test_extract_file_text_txt():
    from app.services.ingestion import extract_file_text

    text = extract_file_text(b"Hello world", "test.txt")
    assert text == "Hello world"


def test_extract_file_text_md():
    from app.services.ingestion import extract_file_text

    md_bytes = b"# Heading\n\nParagraph text."
    text = extract_file_text(md_bytes, "readme.md")
    assert "Heading" in text
    assert "Paragraph text" in text


def test_extract_file_text_txt_with_encoding_errors():
    from app.services.ingestion import extract_file_text

    # Bytes that are not valid UTF-8; should not raise
    bad_bytes = b"\xff\xfe Hello"
    text = extract_file_text(bad_bytes, "doc.txt")
    assert isinstance(text, str)


def test_extract_file_text_unknown_extension_returns_empty():
    from app.services.ingestion import extract_file_text

    text = extract_file_text(b"some binary", "file.pdf")
    assert text == ""


def test_extract_file_text_docx(tmp_path):
    """Test .docx extraction using a real in-memory docx document."""
    pytest.importorskip("docx")
    from docx import Document
    from app.services.ingestion import extract_file_text

    doc = Document()
    doc.add_paragraph("Brand voice paragraph one.")
    doc.add_paragraph("Brand voice paragraph two.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    text = extract_file_text(buf.read(), "brand.docx")
    assert "Brand voice paragraph one." in text
    assert "Brand voice paragraph two." in text


# ── ScrapingError ─────────────────────────────────────────────────────────────

def test_scraping_error_is_exception():
    from app.services.ingestion import ScrapingError

    err = ScrapingError("test error")
    assert isinstance(err, Exception)
    assert str(err) == "test error"


# ── scrape_website (mocked httpx) ────────────────────────────────────────────

@pytest.mark.asyncio
async def test_scrape_website_non_200_raises_scraping_error():
    from app.services.ingestion import ScrapingError, scrape_website

    mock_resp = MagicMock()
    mock_resp.status_code = 404

    with patch("app.services.ingestion.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__ = AsyncMock(return_value=mock_client)
        mock_client.__aexit__ = AsyncMock(return_value=False)
        mock_client.get = AsyncMock(return_value=mock_resp)
        mock_client_cls.return_value = mock_client

        with pytest.raises(ScrapingError) as exc_info:
            await scrape_website("https://example.com")

    assert "HTTP 404" in str(exc_info.value)


@pytest.mark.asyncio
async def test_scrape_website_timeout_raises_scraping_error():
    import httpx
    from app.services.ingestion import ScrapingError, scrape_website

    with patch("app.services.ingestion.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__ = AsyncMock(return_value=mock_client)
        mock_client.__aexit__ = AsyncMock(return_value=False)
        mock_client.get = AsyncMock(side_effect=httpx.TimeoutException("timed out"))
        mock_client_cls.return_value = mock_client

        with pytest.raises(ScrapingError) as exc_info:
            await scrape_website("https://example.com")

    assert "timed out" in str(exc_info.value).lower()


@pytest.mark.asyncio
async def test_scrape_website_returns_text_from_root_and_posts():
    from app.services.ingestion import scrape_website

    root_html = """
    <html><body>
      <main><p>Root page content</p></main>
      <a href="/blog/post-1">Post 1</a>
      <a href="/blog/post-2">Post 2</a>
    </body></html>
    """

    post_html = "<html><body><article><p>Blog post content</p></article></body></html>"

    call_count = 0

    async def fake_get(url, **kwargs):
        nonlocal call_count
        call_count += 1
        resp = MagicMock()
        resp.status_code = 200
        resp.text = root_html if call_count == 1 else post_html
        return resp

    with patch("app.services.ingestion.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__ = AsyncMock(return_value=mock_client)
        mock_client.__aexit__ = AsyncMock(return_value=False)
        mock_client.get = AsyncMock(side_effect=fake_get)
        mock_client_cls.return_value = mock_client

        result = await scrape_website("https://example.com")

    assert "Root page content" in result
    assert "Blog post content" in result


@pytest.mark.asyncio
async def test_scrape_website_caps_at_50k_chars():
    from app.services.ingestion import scrape_website, MAX_TEXT_CHARS

    huge_text = "A" * 100_000
    root_html = f"<html><body><main><p>{huge_text}</p></main></body></html>"

    mock_resp = MagicMock()
    mock_resp.status_code = 200
    mock_resp.text = root_html

    with patch("app.services.ingestion.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__ = AsyncMock(return_value=mock_client)
        mock_client.__aexit__ = AsyncMock(return_value=False)
        mock_client.get = AsyncMock(return_value=mock_resp)
        mock_client_cls.return_value = mock_client

        result = await scrape_website("https://example.com")

    assert len(result) <= MAX_TEXT_CHARS


# ── extract_voice_profile stub ────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_extract_voice_profile_stub_returns_dict():
    import uuid
    from app.services.ingestion import extract_voice_profile

    result = await extract_voice_profile("some text about the brand", uuid.uuid4())
    assert isinstance(result, dict)


# ── extract_clean_text: role attributes ──────────────────────────────────────

def test_extract_clean_text_strips_role_navigation():
    from app.services.ingestion import extract_clean_text

    html = """
    <html><body>
      <div role="navigation">Nav links</div>
      <main><p>Real body text</p></main>
      <div role="banner">Site banner</div>
    </body></html>
    """
    result = extract_clean_text(html)
    assert "Real body text" in result
    assert "Nav links" not in result
    assert "Site banner" not in result


def test_extract_clean_text_extracts_headings():
    from app.services.ingestion import extract_clean_text

    html = """
    <html><body>
      <article>
        <h1>Main Title</h1>
        <h2>Section Heading</h2>
        <h3>Sub-section</h3>
        <p>Body paragraph</p>
      </article>
    </body></html>
    """
    result = extract_clean_text(html)
    assert "Main Title" in result
    assert "Section Heading" in result
    assert "Sub-section" in result
    assert "Body paragraph" in result


# ── scrape_website: ConnectError ──────────────────────────────────────────────

@pytest.mark.asyncio
async def test_scrape_website_connect_error_raises_scraping_error():
    import httpx
    from app.services.ingestion import ScrapingError, scrape_website

    with patch("app.services.ingestion.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__ = AsyncMock(return_value=mock_client)
        mock_client.__aexit__ = AsyncMock(return_value=False)
        mock_client.get = AsyncMock(side_effect=httpx.ConnectError("refused"))
        mock_client_cls.return_value = mock_client

        with pytest.raises(ScrapingError) as exc_info:
            await scrape_website("https://example.com")

    assert "connection refused" in str(exc_info.value).lower()


# ── scrape_website: no blog URLs discovered → root page only ──────────────────

@pytest.mark.asyncio
async def test_scrape_website_no_blog_urls_returns_root_only():
    from app.services.ingestion import scrape_website

    root_html = """
    <html><body>
      <main><p>About us page content only</p></main>
    </body></html>
    """

    mock_resp = MagicMock()
    mock_resp.status_code = 200
    mock_resp.text = root_html

    with patch("app.services.ingestion.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__ = AsyncMock(return_value=mock_client)
        mock_client.__aexit__ = AsyncMock(return_value=False)
        # Only one GET call should happen (no blog posts to fetch)
        mock_client.get = AsyncMock(return_value=mock_resp)
        mock_client_cls.return_value = mock_client

        result = await scrape_website("https://example.com")

    assert "About us page content only" in result
    # Only the root GET was made
    mock_client.get.assert_called_once()


# ── extract_file_text: corrupted docx ────────────────────────────────────────

def test_extract_file_text_corrupted_docx_returns_empty():
    from app.services.ingestion import extract_file_text

    # Pass garbage bytes as a .docx — python-docx should raise internally
    result = extract_file_text(b"not a real docx file content", "corrupt.docx")
    # Should not raise; returns empty string (logged as warning)
    assert result == ""


# ── scrape_website: text from multiple pages joined correctly ─────────────────

@pytest.mark.asyncio
async def test_scrape_website_joins_pages_with_separator():
    from app.services.ingestion import scrape_website

    root_html = """
    <html><body>
      <main><p>Root content</p></main>
      <a href="/blog/post-1">Post 1</a>
    </body></html>
    """
    post_html = "<html><body><article><p>Post content</p></article></body></html>"

    responses = [root_html, post_html]
    call_index = {"n": 0}

    async def fake_get(url, **kwargs):
        resp = MagicMock()
        resp.status_code = 200
        resp.text = responses[min(call_index["n"], len(responses) - 1)]
        call_index["n"] += 1
        return resp

    with patch("app.services.ingestion.httpx.AsyncClient") as mock_client_cls:
        mock_client = AsyncMock()
        mock_client.__aenter__ = AsyncMock(return_value=mock_client)
        mock_client.__aexit__ = AsyncMock(return_value=False)
        mock_client.get = AsyncMock(side_effect=fake_get)
        mock_client_cls.return_value = mock_client

        result = await scrape_website("https://example.com")

    assert "Root content" in result
    assert "Post content" in result
    assert "---" in result  # pages are joined with separator
